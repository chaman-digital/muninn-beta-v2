import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = docx.Document('PERDIDA DE PATRIA POTESTAD IVAN.docx')
    paragraphs = doc.paragraphs

    # Identify start of 2023 facts to purge
    start_purge_idx = -1
    for i, p in enumerate(paragraphs):
        if "15.-" in p.text and "13 de febrero" in p.text and "2023" in p.text:
            start_purge_idx = i
            break
            
    # Identify start of PRUEBAS
    end_purge_idx = -1
    for i, p in enumerate(paragraphs):
        if "P R U E B A S" in p.text:
            end_purge_idx = i
            break

    if start_purge_idx == -1 or end_purge_idx == -1:
        print(f"Error finding boundaries: start={start_purge_idx}, end={end_purge_idx}")
        return

    # Delete paragraphs between start and end (purging the old 2023-2026 timeline and isolated facts)
    for i in range(start_purge_idx, end_purge_idx):
        p = paragraphs[i]
        p._element.getparent().remove(p._element)

    # Now, target_p is the PRUEBAS heading. We insert BEFORE it.
    target_p = paragraphs[end_purge_idx]

    new_facts = [
        ("EVOLUCIÓN CRONOLÓGICA DEL INCUMPLIMIENTO Y ABANDONO (2023-2026)", True, True),
        ("", False, False),
        ("15.- EJE DE SALUD: OMISIÓN TERAPÉUTICA POST-ACCIDENTE Y RIESGO SANITARIO.", True, False),
        ("Derivado del incidente traumático que la menor I.A.L. sufrió en el año 2019, la autoridad médica y ortopédica ordenó enfáticamente el seguimiento mediante fisioterapia para rehabilitar la escoliosis y estabilizar la marcha. Sin embargo, obra en el expediente que la demandada ignoró rotundamente esta directriz, abandonando cualquier continuación de las terapias desde el cese de su emergencia inicial. El escrutinio de la Bitácora de Salud forense [REPORTE_D_Historia_Clinica.md] acredita que no existe inversión temporal o económica de la madre para sanar el daño físico. Más aún, la negligencia es recurrente: tras las estancias recreativas de la progenitora, la menor regresa sistemáticamente con Enfermedades Respiratorias Agudas y bajas dramáticas de Índice de Masa Corporal (IMC), las cuales la demandada ni previene ni atiende médicamente, obligando al Suscrito a remediar.", False, False),
        ("", False, False),
        ("16.- EJE DE EDUCACIÓN: AUSENTISMO COMO CAUSA DE BAJA Y EL RESCATE PATERNO.", True, False),
        ("La demandada ha pretendido encuadrar las bajas académicas argumentando falta de pago o fallas burocráticas; sin embargo, la extracción técnica de los reportes revela que la expulsión forzosa en la Universidad del Valle de México (UVM) y Queen Mary School fue inducida por la madre. Como lo señala la Matriz de Negligencia Escolar [REPORTE_B_Matriz_Negligencia.md], el 78% de las inasistencias injustificadas se concentran de manera contundente en jueves o viernes, empatando matemáticamente con los puentes de descanso o vacacionales en los que la demandada anteponía su esparcimiento [Colegio Queen Mary 2.m4a]. La baja solicitada por el coordinador escolar D. Carlos Alonso López [UVM Campus San Rafael.m4a] constituyó, por ende, el corolario del abandono materno. Mi intervención como padre y mi decisión en la ruta educativa no fue un abandono económico, sino un necesario y justificado acto de rescate frente a esta negligencia.", False, False),
        ("", False, False),
        ("17.- EJE DE FALSEDAD: TRIANGULACIÓN DEL DOLO Y VIOLENCIA VICARIA.", True, False),
        ("La madre ha fabricado denuncias ante la representación social arguyendo 'Retención Ilegal' de la menor (como la Querella CI-FIDCANNA/59/UI-2C/D/03379/09-2023 instaurada el 11 de septiembre de 2023). Es vital ilustrar la temeridad procesal y el dolo sistemático con la Evidencia Satelital y Forense: durante esa precisa ventana donde me acusa falsamente de violencia y sustracción, ella se encontraba voluntariamente vacacionando en pareja en sitios turísticos de Jalisco y Guanajuato. Los metadatos y publicaciones fotográficas intervenidas [Captura de Pantalla 2023-06-27 a la(s) 16.53.23.png] así como sus confesiones extrajudiciales con su hermana 'Guayabita' [IMG_9434.PNG] lo comprueban irrefutablemente [REPORTE_C_Triangulacion_Falsedades.md]. Esta instrumentalización dolosa de las instituciones para separarme de mi hija constituye, de manera indubitable, un acto grave de Violencia Vicaria.", False, False),
        ("", False, False),
        ("18.- PONDERACIÓN JURISPRUDENCIAL Y ALEJAMIENTO JUSTIFICADO (S.C.J.N.).", True, False),
        ("El abandono temporal expuesto en los recesos (ej., el encierro domiciliario sin alimentos de febrero de 2023, documentado en notas de voz como [Nota de voz.m4a] y [13_de_febrero_2023.m4a], donde la menor registró autolesiones tipo 'cutting' e ideación depresiva severa [Es Un Sueño - IO.m4a]) provocó una comprensible ruptura del vínculo tutelar efectivo.", False, False),
        ("Nuestra hija I.A.L. cuenta con la edad de 15 años. En riguroso acatamiento a la jurisprudencia en materia de Autonomía Progresiva estipulada por la Suprema Corte de Justicia de la Nación (SCJN), su nivel de desarrollo le confiere el ejercicio gradual de sus derechos, siendo su voluntad de habitar bajo el resguardo paterno y en total paz una declaración vinculante y no modificable. La reticencia a regresar al ambiente negligente materno califica legalmente como un Alejamiento Justificado por inminente peligro y abandono moral, tipificado en la Ley y en las determinaciones de pérdida de Patria Potestad del Tribunal Supremo.", False, False),
        ("", False, False),
        ("Conforme al Artículo 414, 323 Séptimus Código Civil para la Ciudad de México y 42 de la LGDNNA, ante la omisión continua de cuidados, el Fraude Domiciliar verificado [Video_21_marzo.mp4], y la agresión sistemática hacia el suscrito que contamina a la adolescente, la Suprema Corte mandata invariablemente la Suspensión Definitiva de la Patria Potestad de la promovida.", True, False),
        ("", False, False)
    ]

    for text, bold, title in new_facts:
        new_p = target_p.insert_paragraph_before("")
        if text:
            run = new_p.add_run(text)
            if bold:
                run.bold = True
            if title:
                run.bold = True
                run.underline = True
                new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save('DEMANDA_P_P_P_IVAN_VERSION_BULLETPROOF.docx')
    print("Documento Bulletproof generado correctamente.")

if __name__ == '__main__':
    main()
