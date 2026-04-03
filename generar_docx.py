import docx

def modify_docx():
    doc = docx.Document('PERDIDA DE PATRIA POTESTAD IVAN.docx')
    paragraphs = doc.paragraphs

    # Find the insertion point before P R U E B A S
    insert_idx = -1
    for i, p in enumerate(paragraphs):
        if "Para acreditar los extremos de los hechos narrados" in p.text:
            insert_idx = i
            break
            
    if insert_idx == -1:
        print("No se encontró el punto de inserción, agregando al final del bloque asumiéndolo.")
        insert_idx = len(paragraphs)

    # We need to insert paragraphs. python-docx doesn't easily let you insert a new paragraph *before* another
    # But `paragraph.insert_paragraph_before()` exists!
    target_p = paragraphs[insert_idx]
    
    new_facts = [
        "",
        "24.- Respecto a la Alienación e Interferencia Parental (Art. 323 Séptimus del Código Civil para la CDMX). La demandada se ha conducido sistemáticamente bloqueando la comunicación y generando rechazo. Ello consta en los audios extraídos del celular de resguardo, como en el archivo 'Boleo 62 10.m4a', donde [Boleo 62 10.m4a - 01:15] la demandada expone directrices de bloqueo, transformando la conciencia de la menor con el objeto de impedir sus vínculos conmigo, incurriendo en violencia psicoemocional y manipulación parental punible.",
        "",
        "25.- Respecto a la Obstrucción de Convivencia y Fraude Domiciliar. Se ha constatado la limitación injustificada de las convivencias. Se ha documentado en video (en análisis pericial, correlativo al metadato GPS del Edificio J, Depto 503) que el domicilio real donde habitan ha sido ocultado deliberadamente en el expediente. La Suprema Corte de Justicia de la Nación (SCJN) ha determinado que la limitación injustificada de convivencias afecta el derecho fundamental a los lazos afectivos e idoneidad parental, justificando la privación de patria potestad.",
        "",
        "26.- Respecto a la Negligencia Educativa de la Demandada. Queda documentado el bajo rendimiento y deserción escolar en Queen Mary School y la Universidad del Valle de México ('Colegio Queen Mary.m4a' y 'UVM Campus San Rafael.m4a'). En lugar de asistir, fomenta la deserción y exige delegar injustamente la carga, como se escucha en [UVM Campus San Rafael.m4a - 02:30]. Lo cual constituye un grave incumplimiento de obligaciones de crianza.",
        "",
        "27.- Respecto al Abandono Temerario y Peligro Inminente. Nuestra adolescente de 15 años (I.A.L.) fue dejada repetidamente sola o enviada lejos sin mi autorización (ej., en San Miguel de Allende durante diciembre 2022 y febrero 2023). Esto consta en peritajes y audios como [Nota de voz.m4a - 00:45], donde relata estar sola y atemorizada ante la salida injustificada de su madre de viaje con su nueva pareja. La SCJN establece fehacientemente que este descuido grave es un incumplimiento radical de deberes, lesionando el derecho y estabilidad integral de la menor.",
        "",
        "28.- Respecto a Trabas, Contradicciones de Viaje y la Falsa Querella por Retención. La denuncia formulada el 11 de septiembre de 2023 en mi contra por supuesta retención ilegal es fabricada y dolosa. Existen registros criptográficos ('Captura de Pantalla 2023-06-27 a la(s) 16.53.23.png' representativa de sus redes) donde ella misma exhibe encontrarse vacacionando en pareja en Guadalajara en fechas donde argüía la sustracción, cometiendo claro dolo, falta de lealtad procesal y Violencia Vicaria agravada con las instituciones.",
        "",
        "29.- CRITERIO DE AUTONOMÍA PROGRESIVA Y VINCULATORIEDAD (15 AÑOS). Solicito a Su Señoría dar altísima ponderación y aplicabilidad expresa al principio de Autonomía Progresiva dictado por la SCJN. Toda vez que nuestra menor cuenta con 15 años cumplidos, su narración de los eventos de descuido, así como de su franca voluntad respecto del resguardo con el Suscrito, no se trata solo de un testimonio o escucha de menor usual, sino de la manifestación de un sujeto de derechos cuyo nivel de participación conforma elemento principal y vinculante sobre el nuevo régimen de custodia derivado de la Pérdida solicitada a la Madre.",
        ""
    ]
    
    for text in new_facts:
        target_p.insert_paragraph_before(text)

    # Now add the Cadena de Custodia at the end of the document
    doc.add_page_break()
    p_heading = doc.add_paragraph()
    p_heading.add_run('ANEXO ÚNICO: CADENA DE CUSTODIA DIGITAL (HASH SHA-256)').bold = True
    
    intro_p = doc.add_paragraph(
        "A fin de garantizar la invariabilidad, autenticidad e integridad de la prueba técnica, electrónica y documental aportada y sustraída orgánicamente mediante el sistema de respaldo y bóveda forense, certifico bajo protesta de decir verdad los acuses criptográficos Hash SHA-256 de los medios probatorios referidos en esta incidencia:"
    )
    
    hashes = {
        'Boleo 62 10.m4a': '0fe7cd2cbe3f14e81d9668a8aae3ba83c9bd56090216f44b4094cf1eecc39f14', 
        'Colegio Queen Mary.m4a': '1a1934b717dc7e6350ed93a187b247d67868a8495b575f85cf21e860c6d449ee',
        'UVM Campus San Rafael.m4a': '51d26a01021de41ff6ad8da5c4c66e9659c8e17873d361392f2009b5dcb9c644',
        'Nota de voz.m4a': 'be525ebd245c38b9516b03e934435b18f3ecf317ab2e454495e59339999e5770',
        'Captura de Pantalla 2023-06-27 a la(s) 16.53.23.png': 'd2f5336995a6ee43818314cb1083e4b27be334c3fb8ad6236e1f3cb931b9f845',
        'JUICIO FAMILIAR.pdf': 'cf18c248e2a49eba0be032e9052c7016d60c944b83a9e1148191dbfbc9e0e2f8'
    }
    
    for filename, filehash in hashes.items():
        doc.add_paragraph(f"• Archivo: {filename} \n  Hash/Integridad: {filehash}")

    doc.save('PERDIDA DE PATRIA POTESTAD IVAN.docx')
    print("Modificaci\u00f3n completada con \u00e9xito.")

if __name__ == '__main__':
    modify_docx()
